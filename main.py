import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from agno.agent import Agent
from agno.models.google import Gemini
from agno.models.openai import OpenAIChat
import google.generativeai as genai
import openai
import os
from dotenv import load_dotenv

load_dotenv()
DEFAULT_GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
DEFAULT_OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

st.set_page_config(
    page_title="Social Media Data Analyzer",
    page_icon="📊",
    layout="wide"
)

st.title("Social Media AI Agent 📊")

st.sidebar.header("Configuration")
# Model selection dropdown
model_choice = st.sidebar.selectbox(
    "Select Model",
    ["Gemini", "OpenAI GPT-4o"],
    help="Choose the AI model to use for analysis"
)

# API key input based on selected model
if model_choice == "Gemini":
    api_key = st.sidebar.text_input(
        "Enter your Google API key",
        type="password",
        value=DEFAULT_GOOGLE_API_KEY or "",
        help="Enter your Google API key to use the Gemini model"
    )
    openai_api_key = None
else:  # OpenAI GPT-4o
    api_key = st.sidebar.text_input(
        "Enter your OpenAI API key",
        type="password",
        value=DEFAULT_OPENAI_API_KEY or "",
        help="Enter your OpenAI API key to use the GPT-4o model"
    )
    openai_api_key = api_key
    api_key = None  # Clear Google API key

if not api_key and not openai_api_key:
    st.warning("Please enter the API key for the selected model in the sidebar to proceed.")
else:
    if model_choice == "Gemini" and api_key:
        genai.configure(api_key=api_key)
    elif model_choice == "OpenAI GPT-4o" and openai_api_key:
        openai.api_key = openai_api_key

    @st.cache_resource
    def initialize_agent(model_choice, api_key, openai_api_key):
        if model_choice == "Gemini":
            model = Gemini(id="gemini-2.0-flash", api_key=api_key)
        else:  # OpenAI GPT-4o
            model = OpenAIChat(id="gpt-4o", api_key=openai_api_key)
        
        return Agent(
            name="Social Media Analyzer",
            model=model,
            description="You are an expert Social Media Strategist AI. Your role is to analyze social media data provided from an Excel file and answer user questions with actionable insights and strategic recommendations.",
            instructions=[
                "Carefully analyze the entire provided social media data.",
                "When answering the user's question, refer specifically to the data columns and content.",
                "If engagement metrics (likes, comments, shares, views, etc.) are present, use them to support your analysis of top/bottom performing posts.",
                "Identify key content themes and topics apparent in the data.",
                "Infer the brand's or individual's positioning based on the analyzed content.",
                "Provide strategic insights: what's working, what's not, potential opportunities, or new content angles.",
                "Structure your response in clear, well-organized markdown. Use headings, bullet points, and bold text for emphasis.",
                "Base your analysis *primarily* on the provided Excel data."
            ],
            markdown=True,
        )

    st.write("Upload an Excel file containing social media data.")
    uploaded_file = st.file_uploader(
        "Upload Excel file",
        type=["xlsx", "xls"],
        help="Supported formats: .xlsx, .xls"
    )

    # Display uploaded data
    if uploaded_file is not None:
        if 'df' not in st.session_state:
            try:
                st.session_state.df = pd.read_excel(uploaded_file)
            except Exception as e:
                st.error(f"Error reading Excel file: {e}")
                st.stop()
        df = st.session_state.df
        with st.expander("View Uploaded Data", expanded=False):
            st.dataframe(df)

    # User prompt input
    st.write("Ask questions about your social media data (e.g., 'What are the top-performing posts and why?', 'What themes emerge from the content?')")
    user_prompt = st.text_area(
        "What insights do you want from the data?",
        value='''Key Analytical Outputs Required:

1. Content Themes
    A. Identify top recurring themes and topics
    B. Cluster similar content types (e.g., leadership quotes, industry news, campaign posts)

2. Performance Breakdown
    A. List the top 10 performing posts by engagement rate and total engagement.
        Include insights: Why did these work well? (e.g., timing, emotion, format, subject, CTA)
    B. List the bottom 10 posts by engagement rate.
        Include insights: What might have held them back? (e.g., off-brand, wrong format, poor timing)

3. Positioning Summary
    A. Based on the published content, what image or thought leadership is being projected — intentionally or unintentionally?
    B. If someone only followed this account, what would they think this brand or person stands for?

4. Strategic Insights
    A. Key Learnings: What works well in terms of timing, tone, format, or subject?
    B. Opportunities to Improve: Gaps, missed angles, inconsistent messaging, overused themes
    C. Recommend a sharper or clearer positioning strategy, if necessary
    D. Suggest new content themes or topics that align with audience interests but are not yet covered''',
        placeholder="Enter your question here (e.g., 'Analyze top posts by engagement and explain their success')",
        help="The AI will analyze the data based on your question and provide detailed insights."
    )

    # Analyze button
    if st.button("🔍 Analyze Data", key="analyze_button"):
        if uploaded_file is None:
            st.warning("Please upload an Excel file first.")
        elif not user_prompt:
            st.warning("Please enter a question or prompt to analyze the data.")
        elif not (api_key or openai_api_key):
            st.warning("Please enter the API key for the selected model.")
        else:
            with st.spinner("Analyzing your social media data..."):
                try:
                    df = st.session_state.df
                    data_str = df.to_markdown(index=False)
                    full_prompt = f"Social media data:\n{data_str}\n\nUser question: {user_prompt}"

                    agent = initialize_agent(model_choice, api_key, openai_api_key)
                    response = agent.run(full_prompt)

                    st.subheader("Analysis Results")
                    st.markdown(response.content)

                    doc = Document()
                    doc.add_heading("Social Media Data Analysis", 0)
                    doc.add_paragraph(response.content)
                    output = BytesIO()
                    doc.save(output)
                    output.seek(0)

                    st.download_button(
                        label="📥 Download Analysis as DOCX",
                        data=output,
                        file_name="social_media_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"An error occurred during analysis: {e}")

st.markdown(
    """
    <style>
    .stTextArea textarea {
        height: 250px;
    }
    </style>
    """,
    unsafe_allow_html=True
)
